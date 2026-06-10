import streamlit as st
import random
import string
import io
import unicodedata
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración de la página de Streamlit
st.set_page_config(
    page_title="Generador Pro de Sopa de Letras",
    page_icon="🧩",
    layout="wide"
)

def limpiar_y_normalizar(texto):
    \"\"\"
    Limpia el texto para sopa de letras en español:
    - Pasa a mayúsculas.
    - Quita espacios y caracteres extraños.
    - Elimina acentos (Á->A, É->E, etc.) pero conserva la Ñ de manera segura.
    \"\"\"
    texto = texto.upper().strip()
    # Mapeo explícito para proteger la Ñ/ñ antes de remover diacríticos
    texto = texto.replace('Ñ', '___P_N___')
    
    # Remover acentos de forma estándar
    texto = ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    )
    
    # Restaurar la Ñ
    texto = texto.replace('___P_N___', 'Ñ')
    
    # Filtrar solo caracteres alfabéticos válidos
    return ''.join(c for c in texto if c.isalpha())

def generar_sopa_algoritmo(palabras_originales, tamaño_fijo=None, dificultad="Media"):
    \"\"\"
    Algoritmo avanzado para posicionar palabras en la cuadrícula.
    \"\"\"
    palabras_limpias = []
    mapeo_palabras = []
    
    for p in palabras_originales:
        pl = limpiar_y_normalizar(p)
        if pl:
            palabras_limpias.append(pl)
            mapeo_palabras.append((p.strip(), pl))
            
    if not palabras_limpias:
        return None, None, None, "No se ingresaron palabras válidas."

    # Determinar tamaño de cuadrícula
    max_largo = max(len(pl) for pl in palabras_limpias)
    if tamaño_fijo and tamaño_fijo >= max_largo:
        n = tamaño_fijo
    else:
        total_caracteres = sum(len(pl) for pl in palabras_limpias)
        n = max(max_largo + 3, int(total_caracteres ** 0.5) + 5)
        n = min(max(n, 12), 25) # Límites razonables para impresión (12x12 a 25x25)

    # Inicializar cuadrículas
    grid = [['' for _ in range(n)] for _ in range(n)]
    solucion = [[' ' for _ in range(n)] for _ in range(n)]

    # Definición de direcciones según dificultad
    # (delta_fila, delta_columna)
    if dificultad == "Fácil":
        direcciones = [(0, 1), (1, 0)] # Solo Horizontal derecha y Vertical abajo
    elif dificultad == "Media":
        direcciones = [(0, 1), (1, 0), (1, 1), (-1, 1)] # Horizontal, Vertical y Diagonales hacia adelante
    else: # Difícil
        direcciones = [(0, 1), (1, 0), (1, 1), (-1, 1), (0, -1), (-1, 0), (-1, -1), (1, -1)] # Todas direcciones e invertidas

    palabras_colocadas = []

    for orig, pl in mapeo_palabras:
        colocada = False
        intentos = 0
        
        while not colocada and intentos < 500:
            df, dc = random.choice(direcciones)
            f_inicio = random.randint(0, n - 1)
            c_inicio = random.randint(0, n - 1)
            
            # Verificar límites de la cuadrícula
            if 0 <= f_inicio + df * (len(pl) - 1) < n and 0 <= c_inicio + dc * (len(pl) - 1) < n:
                puede_colocarse = True
                
                # Validar colisiones
                for i, letra in enumerate(pl):
                    f_actual = f_inicio + df * i
                    c_actual = c_inicio + dc * i
                    char_existente = grid[f_actual][c_actual]
                    if char_existente != '' and char_existente != letra:
                        puede_colocarse = False
                        break
                
                if puede_colocarse:
                    for i, letra in enumerate(pl):
                        f_actual = f_inicio + df * i
                        c_actual = c_inicio + dc * i
                        grid[f_actual][c_actual] = letra
                        solucion[f_actual][c_actual] = letra
                    colocada = True
                    palabras_colocadas.append(orig)
            intentos += 1

    # Rellenar espacios vacíos en la sopa de letras
    for f in range(n):
        for c in range(n):
            if grid[f][c] == '':
                grid[f][c] = random.choice(string.ascii_uppercase).replace('Q', 'O') # Reemplazo menor para balancear letras en español
                # Asegurar que si sale una letra aleatoria sea estándar
                if grid[f][c] not in string.ascii_uppercase:
                    grid[f][c] = 'A'
            if solucion[f][c] == ' ':
                solucion[f][c] = '•'

    return grid, solucion, palabras_colocadas, None

def generar_documento_word(grid, solucion, palabras):
    \"\"\"
    Genera un archivo Word (.docx) perfectamente estructurado y listo para impresión.
    \"\"\"
    doc = Document()
    
    # Configurar márgenes óptimos (Estrecho para maximizar espacio de impresión)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Configuración de estilos tipográficos generales
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas' # Fuente monoespaciada ideal para cuadrículas alineadas
    font.size = Pt(13)

    # ================= PÁGINA 1: JUEGO =================
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("SOPA DE LETRAS")
    run_titulo.font.name = 'Arial'
    run_titulo.font.size = Pt(22)
    run_titulo.bold = True
    
    doc.add_paragraph() # Espacio elegante

    # Insertar la cuadrícula alineada de la sopa
    p_sopa = doc.add_paragraph()
    p_sopa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sopa.paragraph_format.line_spacing = 1.25
    
    for fila in grid:
        texto_fila = "  ".join(fila) # Doble espacio entre letras para legibilidad
        run_fila = p_sopa.add_run(texto_fila + "\n")
        run_fila.bold = True
        run_fila.font.size = Pt(14)

    doc.add_paragraph() # Espacio

    # Sección de palabras a buscar
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("PALABRAS A BUSCAR:")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.bold = True

    # Mostrar palabras organizadas en columnas simples o lista limpia
    p_lista = doc.add_paragraph()
    p_lista.paragraph_format.line_spacing = 1.15
    for idx, p in enumerate(sorted(palabras)):
        run_p = p_lista.add_run(f"□ {p.upper():<20}")
        run_p.font.name = 'Arial'
        run_p.font.size = Pt(10)
        # Añadir salto de línea cada 3 palabras para simular columnas estéticas
        if (idx + 1) % 3 == 0:
            p_lista.add_run("\n")

    # ================= PÁGINA 2: SOLUCIÓN =================
    doc.add_page_break()

    titulo_sol = doc.add_paragraph()
    titulo_sol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tsol = titulo_sol.add_run("SOLUCIÓN DE LA SOPA DE LETRAS")
    run_tsol.font.name = 'Arial'
    run_tsol.font.size = Pt(18)
    run_tsol.bold = True
    
    doc.add_paragraph()

    p_sol = doc.add_paragraph()
    p_sol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sol.paragraph_format.line_spacing = 1.25

    for fila in solucion:
        # Reemplazar los puntos de relleno por espacios vacíos para resaltar las respuestas encontradas
        fila_visible = [letra if letra != '•' else ' ' for letra in fila]
        texto_fila_sol = "  ".join(fila_visible)
        run_fsol = p_sol.add_run(texto_fila_sol + "\n")
        run_fsol.bold = True
        run_fsol.font.size = Pt(14)

    # Guardar documento en un buffer de memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- DISEÑO DE LA INTERFAZ WEB EN STREAMLIT ---
st.markdown("<h2 style='text-align: center; color: #1E3A8A;'>🧩 Generador Automático de Sopa de Letras Pro</h2>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #4B5563;'>Diseña sopas de letras personalizadas listas para exportar a Word e imprimir.</p>", unsafe_allow_html=True)
st.hr()

# Uso de columnas para organizar la interfaz
col_izq, col_der = st.columns([1, 2])

with col_izq:
    st.subheader("🛠️ Configuración")
    
    # Área de entrada para las palabras del usuario
    palabras_defecto = "IPERC\nLOTO\nPREVENCION\nSEGURIDAD\nSALUD OCUPACIONAL\nRIESGOS\nINSPECCION\nTRABAJO SEGURO"
    entrada_texto = st.text_area(
        "Ingresa tus palabras o frases (una por línea):",
        value=palabras_defecto,
        height=250,
        help="Cada salto de línea (Enter) se contará como una palabra o frase independiente."
    )
    
    # Parámetros avanzados de la sopa de letras
    dificultad = st.selectbox("Dificultad del juego:", ["Fácil", "Media", "Difícil"], index=1, 
                               help="Fácil (Horizontal/Vertical), Media (+Diagonales), Difícil (+Invertidas)")
    
    tipo_tamaño = st.radio("Tamaño de la cuadrícula:", ["Automático (Recomendado)", "Manual Personalizado"])
    
    tamaño_manual = None
    if tipo_tamaño == "Manual Personalizado":
        tamaño_manual = st.slider("Dimensiones de la cuadrícula (NxN):", min_value=10, max_value=25, value=15)

    ejecutar = st.button("🚀 Generar Estructuras", type="primary", use_container_width=True)

with col_der:
    st.subheader("🖥️ Vista Previa en Tiempo Real")
    
    if ejecutar or 'sopa_grid' in st.session_state:
        # Si se presionó el botón, procesamos los datos actuales
        if ejecutar:
            lineas = [linea.strip() for linea in entrada_texto.split('\n') if linea.strip()]
            
            if not lineas:
                st.error("Por favor, ingresa al menos una palabra o frase válida.")
            else:
                with st.spinner("Construyendo matriz y cruzando caracteres de forma óptima..."):
                    grid, solucion, colocadas, error = generar_sopa_algoritmo(lineas, tamaño_manual, dificultad)
                    
                    if error:
                        st.error(error)
                    else:
                        st.session_state['sopa_grid'] = grid
                        st.session_state['sopa_solucion'] = solucion
                        st.session_state['sopa_colocadas'] = colocadas
        
        # Recuperar datos del estado de sesión si existen
        if 'sopa_grid' in st.session_state:
            grid = st.session_state['sopa_grid']
            solucion = st.session_state['sopa_solucion']
            colocadas = st.session_state['sopa_colocadas']
            
            # Mostrar métricas de éxito
            st.success(f"¡Sopa generada con éxito! Se posicionaron {len(colocadas)} palabras perfectamente.")
            
            # Generación del archivo Word descargable
            bytes_word = generar_documento_word(grid, solucion, colocadas)
            
            st.download_button(
                label="📥 Descargar Documento de Word Listo para Imprimir",
                data=bytes_word,
                file_name="Sopa_De_Letras_Listo_Para_Imprimir.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            # Tabs elegantes para visualizar el tablero y la solución en web
            tab1, tab2 = st.tabs(["🧩 Cuadrícula de Juego", "💡 Matriz de Solución"])
            
            with tab1:
                texto_preview_juego = '\n'.join([' '.join(f) for f in grid])
                st.code(texto_preview_juego, language="text")
                
                st.markdown("**Palabras ocultas confirmadas:**")
                st.write(", ".join([f"• {p}" for p in colocadas]))
                
            with tab2:
                # Reemplazar puntos por espacios para que se resalte visualmente la solución en web
                sol_preview = [[letra if letra != '•' else ' ' for letra in f] for f in solucion]
                texto_preview_sol = '\n'.join([' '.join(f) for f in sol_preview])
                st.code(texto_preview_sol, language="text")
    else:
        st.info("Configura tus palabras en el panel izquierdo y haz clic en 'Generar Estructuras' para visualizar los resultados aquí.")
